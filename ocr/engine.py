"""
Multi-engine OCR module.

Supports:
- PDF text extraction (PyMuPDF / pdfminer)
- Image OCR (pytesseract / PIL)
- DOCX text extraction (python-docx)
- Plain text passthrough
"""

from __future__ import annotations

import os
import io
import time
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class OCRResult:
    text: str
    pages: int = 0
    engine_used: str = "none"
    processing_time_ms: int = 0


class OCREngine:
    """
    Multi-engine OCR processor.

    Automatically selects the best engine based on file type:
    - PDFs -> PyMuPDF (if available) or pdfminer
    - Images -> pytesseract (if available)
    - DOCX -> python-docx
    - TXT -> passthrough
    """

    def __init__(self):
        self._pymupdf = None
        self._tesseract = None
        self._pytesseract = None
        self._pillow = None
        self._try_import_engines()

    def _try_import_engines(self):
        """Try to import optional OCR engines."""
        try:
            import fitz  # PyMuPDF
            self._pymupdf = fitz
        except ImportError:
            pass
        try:
            import pytesseract
            self._pytesseract = pytesseract
        except ImportError:
            pass
        try:
            from PIL import Image
            self._pillow = Image
        except ImportError:
            pass

    def process(self, file_path: str, file_bytes: bytes | None = None) -> OCRResult:
        """
        Process a document and extract text.

        Args:
            file_path: Path to the document file
            file_bytes: Optional raw bytes (if already loaded)

        Returns:
            OCRResult with extracted text and metadata
        """
        t0 = time.time()

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            result = self._process_pdf(file_path, file_bytes)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"):
            result = self._process_image(file_path, file_bytes)
        elif ext == ".docx":
            result = self._process_docx(file_path, file_bytes)
        elif ext == ".txt":
            result = self._process_text(file_path, file_bytes)
        else:
            # Fallback: try image OCR first, then PDF
            if ext in (".png", ".jpg", ".jpeg"):
                result = self._process_image(file_path, file_bytes)
            else:
                result = OCRResult(text="", engine_used="none")

        dur = int((time.time() - t0) * 1000)
        result.processing_time_ms = dur
        return result

    def _process_pdf(self, file_path: str, file_bytes: bytes | None = None) -> OCRResult:
        """Extract text from PDF using PyMuPDF."""
        if self._pymupdf:
            try:
                if file_bytes:
                    doc = self._pymupdf.open(stream=file_bytes, filetype="pdf")
                else:
                    doc = self._pymupdf.open(file_path)
                pages_text = []
                for page in doc:
                    pages_text.append(page.get_text())
                doc.close()
                return OCRResult(
                    text="\n\n".join(pages_text),
                    pages=len(pages_text),
                    engine_used="pymupdf",
                )
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")

        return OCRResult(
            text="(PDF text extraction unavailable - install PyMuPDF: pip install PyMuPDF)",
            engine_used="none",
        )

    def _process_image(self, file_path: str, file_bytes: bytes | None = None) -> OCRResult:
        """OCR on images using pytesseract."""
        if self._pytesseract and self._pillow:
            try:
                if file_bytes:
                    img = self._pillow.open(io.BytesIO(file_bytes))
                else:
                    img = self._pillow.open(file_path)
                text = self._pytesseract.image_to_string(img)
                return OCRResult(text=text.strip(), pages=1, engine_used="tesseract")
            except Exception as e:
                logger.warning(f"Tesseract OCR failed: {e}")

        return OCRResult(
            text="(Image OCR unavailable - install: pip install pytesseract pillow)",
            engine_used="none",
        )

    def _process_docx(self, file_path: str, file_bytes: bytes | None = None) -> OCRResult:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document as DocxDocument
            if file_bytes:
                import io
                doc = DocxDocument(io.BytesIO(file_bytes))
            else:
                doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return OCRResult(
                text="\n".join(paragraphs),
                pages=1,
                engine_used="python-docx",
            )
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            return OCRResult(text="", engine_used="none")

    def _process_text(self, file_path: str, file_bytes: bytes | None = None) -> OCRResult:
        """Passthrough for plain text files."""
        try:
            if file_bytes:
                text = file_bytes.decode("utf-8", errors="replace")
            else:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            return OCRResult(text=text.strip(), pages=1, engine_used="passthrough")
        except Exception as e:
            logger.warning(f"Text extraction failed: {e}")
            return OCRResult(text="", engine_used="none")
